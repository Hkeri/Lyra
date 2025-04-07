import webbrowser
from time import sleep
import pyautogui
from re import search
from keyboard import press_and_release
import os
from psutil import sensors_battery, net_io_counters
from screen_brightness_control import set_brightness
from django.core.mail import send_mail
import ctypes
import os
import psutil
import tempfile
import os
import requests
import speech_recognition as sr
import shutil
from pypdf import PdfReader
import easyocr
import pywhatkit
import datetime
from pytube import YouTube
import qrcode
from groq import Groq
from docx import Document
import time
import keyboard
from googletrans import Translator

def translate_document(input_file, output_file, target_language="en"):
    translator = Translator()
    with open(input_file, 'r') as infile, open(output_file, 'w') as outfile:
        for line in infile:
            translation = translator.translate(line, dest=target_language)
            outfile.write(translation.text + "\n")

translate_document("document.txt", "translated_document.txt", "es")

def s_h():
    webbrowser.open("www.google.com")
    time.sleep(4)
    keyboard.press_and_release('ctrl + h')
    print("Opened Your Search History")

def groq(instrucations, query):
      api_key = "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"
      client = Groq(api_key=api_key)
      chat_completion = client.chat.completions.create(
      messages=[
            {
                  "role": "user",
                  "content": f"{instrucations}, {query}",
            }
      ],
      model="llama-3.3-70b-versatile",
      stream=False,
      )
      ra = (chat_completion.choices[0].message.content).replace("**", "")
      return (ra)

def qrCodeGenerator(input_Text_link):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=15,
        border=4,
    )
    QRfile_name = (str(datetime.now().strftime("%d-%m-%Y"))).replace(" ", "-")
    QRfile_name = QRfile_name.replace(":", "-")
    QRfile_name = QRfile_name.replace(".", "-")
    QRfile_name = QRfile_name + "-QrCode.png"
    qr.add_data(input_Text_link)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")
    img.save(f"{QRfile_name}")
    os.startfile(f"{os.getcwd()}\\{QRfile_name}")

def ytDownloader(yt_url):
    yt = YouTube(yt_url)
    video = yt.streams.get_highest_resolution()
    video.download()
    print("Downloaded the video successfully!")

def whatsapp_send():
    
    mode = input("Contact or Group: ").lower()
    try:
        if mode == "contact":
            phone_number = input("Phone Number (country code as well with +): ")
            message = input(f"What Text Do You Want To Send to {phone_number}: ")

            time_hour = datetime.datetime.now().hour
            time_minute = datetime.datetime.now().minute      
            close_tab = True

            pywhatkit.sendwhatmsg(
                phone_number,
                message,
                time_hour,
                time_minute,
                close_tab,
            )
            print("I Have Sent Your Message")
    except Exception as e:
            print(f"Message Could Not be Sent to {phone_number}: {str(e)}")

    if mode == "group":
        try:
            group_id = input("Please type in the group id: ")
            msg = input("What Text Do You Want To Send To The Group: ")
            time_hour = datetime.datetime.now().hour
            time_minute = datetime.datetime.now().minute
            close_tab = True

            pywhatkit.sendwhatmsg_to_group(
                group_id,
                message,
                time_hour,
                time_minute,
                close_tab,
            )
            print("I Have Sent Your Message to the Group")
        except Exception as e:
            print(f"Message Could Not be Sent to The Group: {str(e)}")

    else:
        print("Please Type in The Following Options: 'Contact' or 'Group'")
        whatsapp_send()

def ocr(image_path):
    # Initialize the EasyOCR reader
    reader = easyocr.Reader(['en'])  # Specify the language(s) you want to use

    # Path to the image file
    image_path = image_path  # Replace with your image path

    # Perform OCR on the image
    results = reader.readtext(image_path)

    # Print the detected text
    for (bbox, text, prob) in results:
        return (text)

def playMusic(song_name):
    try:
        import pywhatkit
        pywhatkit.playonyt(song_name)
        print(f"Playing {song_name}")
    except Exception as e:
        print("Sir I cannot play that certain song")
    
def read_pdf(pdf_file):
    # creating a pdf reader object
    reader = PdfReader(pdf_file)

    # printing number of pages in pdf file
    print(len(reader.pages))

    # getting a specific page from the pdf file
    page = reader.pages[0]

    # extracting text from page
    text = page.extract_text()
    return text

def organize_files(directory):
    os.chdir(directory)

    for filename in os.listdir(directory):
        if os.path.isdir(filename):
            continue
        
        file_extension = filename.split('.')[-1] if '.' in filename else 'no_extension'
        
        # Create specific folders for Word and Excel files
        if file_extension in ['doc', 'docx']:
            target_folder = 'Word'
        elif file_extension in ['xls', 'xlsx']:
            target_folder = 'Excel'
        elif file_extension in ['pptx']:
            target_folder = 'PowerPoint'
        elif file_extension in ['pdf']:
            target_folder = 'PDF'
        elif target_folder in ['exe']:
            target_folder = 'Applications'
        elif file_extension in ['py', 'java', 'html', 'css', 'js']:
            target_folder = 'Code'
        else:
            target_folder = file_extension
        
        if not os.path.exists(target_folder):
            os.makedirs(target_folder)
        
        shutil.move(filename, os.path.join(target_folder, filename))
        print(f'Moved: {filename} to {target_folder}/')

def file_organizer(directory):
    if os.path.isdir(directory):
        organize_files(directory)
        return ("Files organized successfully.")
    else:
        return ("The specified directory does not exist.")

def transcribe_audio(file_path):
    # Create a speech recognition object
    r = sr.Recognizer()

    # Open the audio file
    with sr.AudioFile(file_path) as source:
        # Read the audio data
        audio = r.record(source)

        # Transcribe the audio
        try:
            transcription = r.recognize_google(audio)
            return transcription
        except sr.UnknownValueError:
            return "Could not understand audio"
        except sr.RequestError:
            return "Could not request results"

def download_images(image_urls):
    for url in image_urls:
        response = requests.get(url)
        if response.status_code == 200:
            image_name = url.split("/")[-1]
            with open(image_name, 'wb') as img_file:
                img_file.write(response.content)
            return (f"Downloaded: {image_name}")
        else:
             return (f"Failed to download {url}: {response.status_code}")

def clean_temp():
    temp_dir = tempfile.gettempdir()
    for filename in os.listdir(temp_dir):
        file_path = os.path.join(temp_dir, filename)
        try:
            os.remove(file_path)
        except Exception as e:
            return (f"Error deleting file {file_path}: {e}")

def get_file_extension(text):
    text = text.lower()
    if "python file" in text:
        ex = ".py"
    elif "java file" in text:
        ex = ".java"
    elif "text file" in text:
        ex = ".txt"
    elif "html file" in text:
        ex = ".html"
    elif "css file" in text:
        ex = ".css"
    elif "javascript file" in text:
        ex = ".js"
    elif "json file" in text:
        ex = ".json"
    elif "xml file" in text:
        ex = ".xml"
    elif "csv file" in text:
        ex = ".csv"
    elif "markdown file" in text:
        ex = ".md"
    elif "yaml file" in text:
        ex = ".yaml"
    elif "pdf file" in text:
        ex = ".pdf"
    elif "word file" in text:
        ex = ".docx"
    elif "excel file" in text:
        ex = ".xlsx"
    elif "powerpoint file" in text:
        ex = ".pptx"
    elif "zip file" in text:
        ex = ".zip"
    elif "tar file" in text:
        ex = ".tar"
    else:
        ex = ""  # Default case if no match found
    return ex

def update_text(text):
    if "python file" in text:
        text = text.replace("python file", "")
    elif "java file" in text:
        text = text.replace("java file", "")
    elif "text file" in text:
        text = text.replace("text file", "")
    elif "html file" in text:
        text = text.replace("html file", "")
    elif "css file" in text:
        text = text.replace("css file", "")
    elif "javascript file" in text:
        text = text.replace("javascript file", "")
    elif "json file" in text:
        text = text.replace("json file", "")
    elif "xml file" in text:
        text = text.replace("xml file", "")
    elif "csv file" in text:
        text = text.replace("csv file", "")
    elif "markdown file" in text:
        text = text.replace("markdown file", "")
    elif "yaml file" in text:
        text = text.replace("yaml file", "")
    elif "image file" in text:
        text = text.replace("image file", "")
    elif "video file" in text:
        text = text.replace("video file", "")
    elif "audio file" in text:
        text = text.replace("audio file", "")
    elif "pdf file" in text:
        text = text.replace("pdf file", "")
    elif "word file" in text:
        text = text.replace("word file", "")
    elif "excel file" in text:
        text = text.replace("excel file", "")
    elif "powerpoint file" in text:
        text = text.replace("powerpoint file", "")
    elif "zip file" in text:
        text = text.replace("zip file", "")
    elif "tar file" in text:
        text = text.replace("tar file", "")
    else:
        pass
    return text



def create_file(text):
    selected_ex = get_file_extension(text)
    text = update_text(text)
    if "named" in text or "with name" in text:
        text = text.replace("named","")
        text = text.replace("with name","")
        text = text.replace("create","")
        text = text.strip()
        with open(f"{text}{selected_ex}","w"):
            pass
    else :
        with open(f"demo{selected_ex}","w"):
            pass

def get_top_processes(num_processes=3):
    # Create a list to hold process information
    processes = []

    # Iterate through all running processes
    for process in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_info']):
        try:
            # Get CPU and memory usage
            cpu_usage = process.info['cpu_percent']
            memory_usage = process.info['memory_info'].rss  # Resident Set Size (RSS)

            # Append process info to the list
            processes.append((process.info['name'], cpu_usage, memory_usage))
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            # Handle the case where the process has terminated or access is denied
            continue

    # Sort processes by CPU usage and then by memory usage
    processes.sort(key=lambda p: (p[1], p[2]), reverse=True)

    # Get the top 'num_processes' processes
    top_processes = processes[:num_processes]

    return top_processes

def display_top_processes():
    top_processes = get_top_processes()
    for name, cpu, memory in top_processes:
        return (f"Process: {name}, CPU Usage: {int(cpu)}, Memory Usage: {int(memory / (1024 * 1024))}")  # Convert bytes to MB
    

def change_wallpaper(image_path):
    # Check if the file exists
    if not os.path.isfile(image_path):
        print(f"The file {image_path} does not exist.")
        print("You Can Go back to The main Menu")

    # Use the SystemParametersInfo function from user32.dll
    try:
        # SPI_SETDESKWALLPAPER = 20
        # SPIF_UPDATEINIFILE = 0x01
        # SPIF_SENDCHANGE = 0x02
        ctypes.windll.user32.SystemParametersInfoW(20, 0, image_path, 3)
        print("Wallpaper changed successfully.")
    except Exception as e:
        print(f"Failed to change wallpaper: {e}")

def analyze_and_report(csv_file, report_file):
    """Analyzes data and generates a report using Groq AI."""

    try:
        with open(csv_file, 'r') as file:
            csv_content = file.read()

        # Use Groq AI to generate the report
        instructions = "Analyze the following CSV data and generate a detailed report in Word format"
        query = f"CSV Data:\n{csv_content}"
        report_content = groq(instructions, query)

        # Save the report content to a Word document
        document = Document()
        document.add_heading("AI-Generated Report", level=1)
        document.add_paragraph(report_content)
        document.save(report_file)

        print(f"Report generated in '{report_file}'.")

    except FileNotFoundError:
        print(f"Error: File '{csv_file}' not found.")
    except Exception as e:
        print(f"An error occurred: {e}")


def send_email(message, email):
    send_mail(
        "",
        message,
        email,
        fail_silently=False,
    )

def dim_light():
    set_brightness(45)

def internet_speed():
        # Select network interface (replace 'en0' with your interface)
        net_if = net_io_counters(pernic=True)['en0']

        # Initial byte counts
        bytes_recv = net_if.bytes_recv

        sleep(3)  # Measure for 3 seconds

        # Calculate bytes transferred
        new_bytes_recv = net_if.bytes_recv
        recv = new_bytes_recv - bytes_recv

        # Convert to Mbps
        download_speed = recv / (5 * 1024 * 1024)

        return f"Internet Speed: {download_speed:.2f} Mbps"

def smart_battery():
    batt = sensors_battery()
    if batt.power_plugged:
        print(f"The Battery is Currently {batt.percent}")

    else:
        if batt.percent <= 75:
            print(
                f"The Battery is {batt.percent} and it is Perfect!"
            )

        elif batt.percent <= 50 and batt.percent >= 75:
            print(
                f"The Battery is {batt.percent} and it is at good charge!"
            )

        elif batt.percent <= 25 and batt.percent >= 50:
            print(
                f"The Battery is {batt.percent} but you \nhave to charge since it is kinda low!"
            )

        elif batt.percent <= 10 and batt.percent >= 25:
            print(
                f"The Battery is {batt.percent} so really you must charge it right away!"
            )

        elif batt.percent <= 5 and batt.percent >= 10:
            print(
                f"Sir, You must charge you computer right now because it is {batt.percent}"
            )

        else:
            print(
                f"Sir it is Extremely Low because it is {batt.percent}. \nCharge it Immediately!"
            )


def yt_search(user):
    user = user.replace("youtube search", "")
    user = user.replace("Youtube search", "")
    user = user.replace("youtube Search", "")
    webbrowser.open(f"https://www.youtube.com/results?search_query={user}")
    return (
        "Sir! This is what I found on Youtube According to Your user!"
    )

def openappweb(query):
    if search(".com" or ".co" or ".org", query):
        query = query.replace("open", "")
        query = query.replace("jarvis", "")
        query = query.replace("launch", "")
        query = query.replace(" ", "")
        webbrowser.open(f"https://www.{query}")
    else:
        file_to_open = query.lower()
        file_to_open = query.replace("open ", "")
        file_to_open = file_to_open.replace("launch", "")
        file_to_open = file_to_open.replace("start", "")
        file_to_open = file_to_open.replace("command prompt", "cmd")
        file_to_open = file_to_open.replace("word", "winword")
        file_to_open = file_to_open.replace("vscode", "code")
        file_to_open = file_to_open.replace("visual studio code", "code")
        file_to_open = file_to_open.replace("power point", "powerpnt")
        file_to_open = file_to_open.replace("google chrome", "chrome")
        file_to_open = file_to_open.replace("edge", "msedge")
        file_to_open = file_to_open.replace("calculator", "calc")
        file_to_open = file_to_open.replace("paint", "mspaint")
        file_to_open = file_to_open.replace("file explorer", "explorer")
        file_to_open = file_to_open.replace("task manager", "taskmgr")
        file_to_open = file_to_open.replace("settings", "ms-settings:")
        file_to_open = file_to_open.replace("control panel", "control")
        file_to_open = file_to_open.replace("camera", "microsoft.windows.camera:")
        file_to_open = file_to_open.replace("store", "ms-windows-store:")
        file_to_open = file_to_open.replace("snipping tool", "snippingtool")
        file_to_open = file_to_open.replace("regedit", "regedit")
        file_to_open = file_to_open.replace("msconfig", "msconfig")
        file_to_open = file_to_open.replace("device manager", "devmgmt.msc")
        file_to_open = file_to_open.replace("disk manager", "diskmgmt.msc")
        file_to_open = file_to_open.replace("event viewer", "eventvwr.msc")
        file_to_open = file_to_open.replace("services", "services.msc")
        file_to_open = file_to_open.replace("system information", "msinfo32")
        file_to_open = file_to_open.replace("performance monitor", "perfmon.msc")
        file_to_open = file_to_open.replace("resource monitor", "resmon.exe")
        file_to_open = file_to_open.replace("character map", "charmap")
        file_to_open = file_to_open.replace("magnifier", "magnify")
        file_to_open = file_to_open.replace("remote desktop connection", "mstsc")
        file_to_open = file_to_open.replace("powershell ise", "powershell_ise")
        file_to_open = file_to_open.replace("paint 3d", "ms-paint:")
        file_to_open = file_to_open.replace("directx diagnostic tool", "dxdiag")
        file_to_open = file_to_open.replace("font viewer", "fontview")
        file_to_open = file_to_open.replace("iexpress wizard", "iexpress")
        file_to_open = file_to_open.replace("lpksetup", "lpksetup")
        file_to_open = file_to_open.replace("mmc", "mmc")
        file_to_open = file_to_open.replace("microsoft remote assistance", "msra")
        file_to_open = file_to_open.replace("narrator", "narrator")
        file_to_open = file_to_open.replace("on-screen keyboard", "osk")
        file_to_open = file_to_open.replace("print management", "printmanagement.msc")
        file_to_open = file_to_open.replace("phone and modem options", "telephon.cpl")
        file_to_open = file_to_open.replace("date and time", "timedate.cpl")
        file_to_open = file_to_open.replace("utility manager", "utilman")
        file_to_open = file_to_open.replace("windows image acquisition wizard", "wiaacmgr")
        file_to_open = file_to_open.replace("windows version", "winver")
        file_to_open = file_to_open.replace("ip configuration", "ipconfig")
        file_to_open = file_to_open.replace("trace route", "tracert")
        file_to_open = file_to_open.replace("network statistics", "netstat")
        file_to_open = file_to_open.replace("system file checker", "sfc")
        file_to_open = file_to_open.replace("open gl viewer", "glview")
        file_to_open = file_to_open.replace("ms access", "msaccess")
        file_to_open = file_to_open.replace("disk cleanup", "cleanmgr")
        file_to_open = file_to_open.replace("reliability monitor", "perfmon /rel")
        file_to_open = file_to_open.replace("group policy editor", "gpedit.msc")
        file_to_open = file_to_open.replace("local security policy", "secpol.msc")
        file_to_open = file_to_open.replace("computer management", "compmgmt.msc")
        file_to_open = file_to_open.replace("odbc data source administrator", "odbcad32")
        file_to_open = file_to_open.replace("system configuration editor", "sysedit")
        file_to_open = file_to_open.replace("task scheduler", "taskschd.msc")
        file_to_open = file_to_open.replace("windows defender firewall", "wf.msc")
        file_to_open = file_to_open.replace("wordpad", "wordpad")
        file_to_open = file_to_open.replace("ping", "ping")
        file_to_open = file_to_open.replace("storage spaces", "storagespaces")
        file_to_open = file_to_open.replace("ms publisher", "mspub")
        file_to_open = file_to_open.replace("windows fax and scan", "wfs.exe")
        file_to_open = file_to_open.replace("ie4uinit", "ie4uinit") #Internet explorer per-user initialization utility.
        file_to_open = file_to_open.replace("iscsicpl", "iscsicpl") #iSCSI initiator.
        file_to_open = file_to_open.replace("journal", "journal") #Windows journal.
        file_to_open = file_to_open.replace("lodctr", "lodctr") #Performance counter.
        file_to_open = file_to_open.replace("mobsync", "mobsync") #Synchronization Manager
        file_to_open = file_to_open.replace("msdt", "msdt") #Microsoft support diagnostic tool.
        file_to_open = file_to_open.replace("msiexec", "msiexec") #Windows installer.
        file_to_open = file_to_open.replace("ntbackup", "ntbackup") #Windows backup.
        file_to_open = file_to_open.replace("ocsetup", "ocsetup") #Windows optional component setup.
        file_to_open = file_to_open.replace("printui", "printui") #Print user interface.
        file_to_open = file_to_open.replace("rasdial", "rasdial") #Remote access dialer.
        file_to_open = file_to_open.replace("route", "route") #Network routing.
        file_to_open = file_to_open.replace("shrpubw", "shrpubw") #Share folder wizard.
        file_to_open = file_to_open.replace("slui", "slui") #Windows activation.
        file_to_open = file_to_open.replace("subst", "subst") #Substitutes a path.
        file_to_open = file_to_open.replace("syncapp", "syncapp.cpl") #Sync center.
        file_to_open = file_to_open.replace("takeown", "takeown") #Take ownership of a file.
        file_to_open = file_to_open.replace("tpm.msc", "tpm.msc") #TPM management.
        file_to_open = file_to_open.replace("tsdiscon", "tsdiscon") #Disconnect remote desktop.
        file_to_open = file_to_open.replace("tskill", "tskill") #Kill a remote desktop process.
        file_to_open = file_to_open.replace("verifier", "verifier") #Driver verifier.
        file_to_open = file_to_open.replace("wab", "wab") #Windows address book.
        file_to_open = file_to_open.replace("wermgr", "wermgr") #Windows error reporting.
        file_to_open = file_to_open.replace("winchat", "winchat") #Windows chat.
        file_to_open = file_to_open.replace("wmic", "wmic") #Windows management instrumentation command-line.
        file_to_open = file_to_open.replace("mshtmlhelp", "hh.exe") #HTML help.
        file_to_open = file_to_open.replace("wiavideo", "wiavideo.exe")
        press_and_release("win + r")
        sleep(0.5)
        pyautogui.typewrite(f"{file_to_open.lower()}.exe")
        sleep(0.5)
        press_and_release("enter")

def closeappweb(query):
    if search("tab", query):
        pyautogui.hotkey("ctrl", "w")
    else:
        file_to_close = query.lower()
        file_to_close = query.replace("stop", "")
        file_to_close = query.replace("close ", "")
        file_to_close = query.replace("exit ", "")
        file_to_close = file_to_close.replace("launch", "")
        file_to_open = file_to_open.replace("command prompt", "cmd")
        file_to_open = file_to_open.replace("word", "winword")
        file_to_open = file_to_open.replace("vscode", "code")
        file_to_open = file_to_open.replace("visual studio code", "code")
        file_to_open = file_to_open.replace("power point", "powerpnt")
        file_to_open = file_to_open.replace("google chrome", "chrome")
        file_to_open = file_to_open.replace("edge", "msedge")
        file_to_open = file_to_open.replace("calculator", "calc")
        file_to_open = file_to_open.replace("paint", "mspaint")
        file_to_open = file_to_open.replace("file explorer", "explorer")
        file_to_open = file_to_open.replace("task manager", "taskmgr")
        file_to_open = file_to_open.replace("settings", "ms-settings:")
        file_to_open = file_to_open.replace("control panel", "control")
        file_to_open = file_to_open.replace("camera", "microsoft.windows.camera:")
        file_to_open = file_to_open.replace("store", "ms-windows-store:")
        file_to_open = file_to_open.replace("snipping tool", "snippingtool")
        file_to_open = file_to_open.replace("regedit", "regedit")
        file_to_open = file_to_open.replace("msconfig", "msconfig")
        file_to_open = file_to_open.replace("device manager", "devmgmt.msc")
        file_to_open = file_to_open.replace("disk manager", "diskmgmt.msc")
        file_to_open = file_to_open.replace("event viewer", "eventvwr.msc")
        file_to_open = file_to_open.replace("services", "services.msc")
        file_to_open = file_to_open.replace("system information", "msinfo32")
        file_to_open = file_to_open.replace("performance monitor", "perfmon.msc")
        file_to_open = file_to_open.replace("resource monitor", "resmon.exe")
        file_to_open = file_to_open.replace("character map", "charmap")
        file_to_open = file_to_open.replace("magnifier", "magnify")
        file_to_open = file_to_open.replace("remote desktop connection", "mstsc")
        file_to_open = file_to_open.replace("powershell ise", "powershell_ise")
        file_to_open = file_to_open.replace("paint 3d", "ms-paint:")
        file_to_open = file_to_open.replace("directx diagnostic tool", "dxdiag")
        file_to_open = file_to_open.replace("font viewer", "fontview")
        file_to_open = file_to_open.replace("iexpress wizard", "iexpress")
        file_to_open = file_to_open.replace("lpksetup", "lpksetup")
        file_to_open = file_to_open.replace("mmc", "mmc")
        file_to_open = file_to_open.replace("microsoft remote assistance", "msra")
        file_to_open = file_to_open.replace("narrator", "narrator")
        file_to_open = file_to_open.replace("on-screen keyboard", "osk")
        file_to_open = file_to_open.replace("print management", "printmanagement.msc")
        file_to_open = file_to_open.replace("phone and modem options", "telephon.cpl")
        file_to_open = file_to_open.replace("date and time", "timedate.cpl")
        file_to_open = file_to_open.replace("utility manager", "utilman")
        file_to_open = file_to_open.replace("windows image acquisition wizard", "wiaacmgr")
        file_to_open = file_to_open.replace("windows version", "winver")
        file_to_open = file_to_open.replace("ip configuration", "ipconfig")
        file_to_open = file_to_open.replace("trace route", "tracert")
        file_to_open = file_to_open.replace("network statistics", "netstat")
        file_to_open = file_to_open.replace("system file checker", "sfc")
        file_to_open = file_to_open.replace("open gl viewer", "glview")
        file_to_open = file_to_open.replace("ms access", "msaccess")
        file_to_open = file_to_open.replace("disk cleanup", "cleanmgr")
        file_to_open = file_to_open.replace("reliability monitor", "perfmon /rel")
        file_to_open = file_to_open.replace("group policy editor", "gpedit.msc")
        file_to_open = file_to_open.replace("local security policy", "secpol.msc")
        file_to_open = file_to_open.replace("computer management", "compmgmt.msc")
        file_to_open = file_to_open.replace("odbc data source administrator", "odbcad32")
        file_to_open = file_to_open.replace("system configuration editor", "sysedit")
        file_to_open = file_to_open.replace("task scheduler", "taskschd.msc")
        file_to_open = file_to_open.replace("windows defender firewall", "wf.msc")
        file_to_open = file_to_open.replace("wordpad", "wordpad")
        file_to_open = file_to_open.replace("ping", "ping")
        file_to_open = file_to_open.replace("storage spaces", "storagespaces")
        file_to_open = file_to_open.replace("ms publisher", "mspub")
        file_to_open = file_to_open.replace("windows fax and scan", "wfs.exe")
        file_to_open = file_to_open.replace("ie4uinit", "ie4uinit") #Internet explorer per-user initialization utility.
        file_to_open = file_to_open.replace("iscsicpl", "iscsicpl") #iSCSI initiator.
        file_to_open = file_to_open.replace("journal", "journal") #Windows journal.
        file_to_open = file_to_open.replace("lodctr", "lodctr") #Performance counter.
        file_to_open = file_to_open.replace("mobsync", "mobsync") #Synchronization Manager
        file_to_open = file_to_open.replace("msdt", "msdt") #Microsoft support diagnostic tool.
        file_to_open = file_to_open.replace("msiexec", "msiexec") #Windows installer.
        file_to_open = file_to_open.replace("ntbackup", "ntbackup") #Windows backup.
        file_to_open = file_to_open.replace("ocsetup", "ocsetup") #Windows optional component setup.
        file_to_open = file_to_open.replace("printui", "printui") #Print user interface.
        file_to_open = file_to_open.replace("rasdial", "rasdial") #Remote access dialer.
        file_to_open = file_to_open.replace("route", "route") #Network routing.
        file_to_open = file_to_open.replace("shrpubw", "shrpubw") #Share folder wizard.
        file_to_open = file_to_open.replace("slui", "slui") #Windows activation.
        file_to_open = file_to_open.replace("subst", "subst") #Substitutes a path.
        file_to_open = file_to_open.replace("syncapp", "syncapp.cpl") #Sync center.
        file_to_open = file_to_open.replace("takeown", "takeown") #Take ownership of a file.
        file_to_open = file_to_open.replace("tpm.msc", "tpm.msc") #TPM management.
        file_to_open = file_to_open.replace("tsdiscon", "tsdiscon") #Disconnect remote desktop.
        file_to_open = file_to_open.replace("tskill", "tskill") #Kill a remote desktop process.
        file_to_open = file_to_open.replace("verifier", "verifier") #Driver verifier.
        file_to_open = file_to_open.replace("wab", "wab") #Windows address book.
        file_to_open = file_to_open.replace("wermgr", "wermgr") #Windows error reporting.
        file_to_open = file_to_open.replace("winchat", "winchat") #Windows chat.
        file_to_open = file_to_open.replace("wmic", "wmic") #Windows management instrumentation command-line.
        file_to_open = file_to_open.replace("mshtmlhelp", "hh.exe") #HTML help.
        file_to_open = file_to_open.replace("wiavideo", "wiavideo.exe")
        os.system(f"taskkill /f /im {file_to_close}.exe")
